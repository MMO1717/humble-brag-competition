from copy import deepcopy
from pathlib import Path
from zipfile import ZipFile, ZIP_DEFLATED

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


REF = Path("/Users/mm/Desktop/CST2403_202411119_栾震.docx")
SRC = Path("/Users/mm/Desktop/Bragging_acl2025-main/汇报/项目进度汇报.docx")
OUT = Path("/Users/mm/Desktop/Bragging_acl2025-main/汇报/项目进度汇报_标准格式.docx")


def clear_body_keep_section(doc: Document) -> None:
    body = doc._body._element
    sect_pr = None
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            sect_pr = deepcopy(child)
        body.remove(child)
    if sect_pr is not None:
        body.append(sect_pr)


def set_run_font(run, east="宋体", west="Times New Roman", size=None, bold=None):
    run.font.name = west
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east)
    r_fonts.set(qn("w:ascii"), west)
    r_fonts.set(qn("w:hAnsi"), west)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def add_center_para(doc, text="", size=None, bold=False, style="Normal", before_blank=0):
    for _ in range(before_blank):
        doc.add_paragraph("")
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def add_body_para(doc, text, style="论文正文"):
    parts = text.split("\n")
    first = True
    for part in parts:
        if not part.strip():
            doc.add_paragraph("", style=style)
            continue
        p = doc.add_paragraph(style=style)
        r = p.add_run(part.strip())
        set_run_font(r)
        if first and "：" in part and len(part.split("：", 1)[0]) <= 20:
            # Keep source labels readable without changing the inherited paragraph style.
            p.clear()
            label, rest = part.split("：", 1)
            r1 = p.add_run(label + "：")
            set_run_font(r1, east="黑体", west="Times New Roman", bold=True)
            r2 = p.add_run(rest.strip())
            set_run_font(r2)
        first = False


def add_heading(doc, text, level):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text.strip())
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(r, east="黑体", west="Times New Roman", size=16, bold=True)
    elif level == 2:
        set_run_font(r, east="黑体", west="Times New Roman", size=15, bold=True)
    else:
        set_run_font(r, east="黑体", west="Times New Roman", size=14, bold=True)
    return p


def add_blank_image_box(doc, height_cm):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(13.0)
    cell = table.cell(0, 0)
    cell.width = Cm(13.0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    row = table.rows[0]
    row.height = Cm(height_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.tcW
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(13.0).twips)))
    tc_w.set(qn("w:type"), "dxa")
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")
        borders.append(el)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("", style="论文正文")


def copy_report_content(doc):
    src = Document(SRC)
    for i, p in enumerate(src.paragraphs):
        text = p.text.strip()
        has_image = bool(p._p.xpath(".//w:drawing"))
        if text:
            if p.style.name == "Heading 1":
                add_heading(doc, text, 1)
            elif p.style.name == "Heading 2":
                add_heading(doc, text, 2)
            else:
                add_body_para(doc, text)
        if has_image:
            add_blank_image_box(doc, 6.8 if i == 10 else 2.0)


def main():
    doc = Document(REF)
    clear_body_keep_section(doc)

    add_center_para(doc, "《项目进度汇报》", size=17.5, bold=True, before_blank=5)
    add_center_para(doc, "基于 LLM 的低调炫耀理解与回应生成", size=15, bold=True, before_blank=1)
    add_center_para(doc, "系统设计与阶段性测试", size=15, bold=True)
    add_center_para(doc, "东北大学", size=10.5, bold=True, before_blank=10)
    add_center_para(doc, "2026年05月", size=10.5, bold=True)

    doc.add_page_break()
    add_center_para(doc, "基于 LLM 的低调炫耀理解与回应生成系统项目进度汇报", style="论文正文")
    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("摘    要")
    set_run_font(r, east="黑体", west="Times New Roman", size=16, bold=True)
    add_body_para(
        doc,
        "本文围绕低调炫耀（Humble Bragging）理解与回应生成任务，汇总了当前项目在数据结构分析、任务语义理解、智能体架构设计和阶段性评测方面的进展。系统以社交语境、发帖人意图、期望反馈和回应策略为核心约束，通过模块化提示词、结构化输出校验和条件式自纠正机制，提升模型对炫耀机制识别、策略选择和自然回复生成的稳定性。",
    )
    add_body_para(doc, "关键词：低调炫耀；大语言模型；回应生成；策略选择；自纠正")

    copy_report_content(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    strip_unreferenced_media(OUT)
    print(OUT)


def strip_unreferenced_media(path: Path) -> None:
    tmp = path.with_suffix(".cleaning.docx")
    with ZipFile(path, "r") as zin, ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            name = item.filename
            if name.startswith("word/media/"):
                continue
            data = zin.read(name)
            if name.endswith(".rels"):
                text = data.decode("utf-8")
                import re

                text = re.sub(
                    r'<Relationship[^>]+Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"[^>]*/>',
                    "",
                    text,
                )
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


if __name__ == "__main__":
    main()
